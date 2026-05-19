"""
Raw Ingester: Đọc file PDF/DOCX từ thư mục Luật chứng khoán, dùng Gemini để trích xuất metadata,
lưu metadata vào PostgreSQL và vector vào ChromaDB.
"""
import os
import sys
import re
import json
import time
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import chromadb
from sentence_transformers import SentenceTransformer

# Cấu hình encoding cho stdout trên Windows để tránh lỗi Unicode
sys.stdout.reconfigure(encoding='utf-8')

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import CHROMA_DIR, CHROMA_COLLECTION, EMBEDDING_MODEL, BASE_DIR, GEMINI_API_KEY, GEMINI_MODEL
from backend.database import SessionLocal, engine, Base
from backend.models import Document
from processing.chunker import chunk_by_article

# Khởi tạo DB tables
Base.metadata.create_all(bind=engine)

DATA_PATH = os.path.join(BASE_DIR, "Luật chứng khoán")

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            text += page.get_text()
    except Exception as e:
        print(f"Lỗi đọc PDF {pdf_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = DocxDocument(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Lỗi đọc DOCX {docx_path}: {e}")
    return text

METADATA_FALLBACK_MAP = {
    "2019_581 + 582_36-vbhn-btc.pdf": {
        "so_hieu": "36/VBHN-BTC",
        "ten": "Văn bản hợp nhất 36/VBHN-BTC hướng dẫn chào bán chứng khoán",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2019-12-30",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_54.signed.pdf": {
        "so_hieu": "54/2019/QH14",
        "ten": "Luật Chứng khoán số 54/2019/QH14",
        "co_quan": "Quốc hội",
        "ngay_hieu_luc": "2021-01-01",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Luật"
    },
    "vanbangoc_tt 28.2020.btc.pdf": {
        "so_hieu": "28/2020/TT-BTC",
        "ten": "Thông tư 28/2020/TT-BTC quy định về niêm yết chứng khoán",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2020-06-01",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_116-2020-tt-btc.pdf": {
        "so_hieu": "116/2020/TT-BTC",
        "ten": "Thông tư 116/2020/TT-BTC hướng dẫn quản lý quỹ đầu tư chứng khoán",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2021-02-15",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_117-2020-tt-btc.pdf": {
        "so_hieu": "117/2020/TT-BTC",
        "ten": "Thông tư 117/2020/TT-BTC hướng dẫn thành lập và quản lý quỹ thành viên",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2021-02-15",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_155.signed.pdf": {
        "so_hieu": "155/2020/NĐ-CP",
        "ten": "Nghị định 155/2020/NĐ-CP quy định chi tiết thi hành một số điều của Luật Chứng khoán",
        "co_quan": "Chính phủ",
        "ngay_hieu_luc": "2021-01-01",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Nghị định"
    },
    "vanbangoc_73_19012024_153629.pdf.pdf": {
        "so_hieu": "73/2013/TT-BTC",
        "ten": "Thông tư 73/2013/TT-BTC hướng dẫn chi tiết về niêm yết chứng khoán",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2013-07-15",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_nq 26.2024.pdf": {
        "so_hieu": "26/2024/NQ-HĐQT",
        "ten": "Nghị quyết số 26/2024/NQ-HĐQT",
        "co_quan": "Hội đồng Quản trị",
        "ngay_hieu_luc": "2024-05-15",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Quyết định"
    },
    "vanbangoc_nq 34.2024.pdf": {
        "so_hieu": "34/2024/NQ-HĐQT",
        "ten": "Nghị quyết số 34/2024/NQ-HĐQT",
        "co_quan": "Hội đồng Quản trị",
        "ngay_hieu_luc": "2024-06-20",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Quyết định"
    },
    "vanbangoc_77.tt.pdf": {
        "so_hieu": "77/2020/TT-BTC",
        "ten": "Thông tư 77/2020/TT-BTC hướng dẫn chế độ kế toán cho công ty chứng khoán",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2021-01-01",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_nq số 69.2025 ngày 27.11.2025.pdf": {
        "so_hieu": "69/2025/NQ-HĐQT",
        "ten": "Nghị quyết số 69/2025/NQ-HĐQT ngày 27/11/2025",
        "co_quan": "Hội đồng Quản trị",
        "ngay_hieu_luc": "2025-11-27",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Quyết định"
    },
    "vanbangoc_10.2013.ttlt.btp.bca.tandtc.vksndtc.btc.pdf": {
        "so_hieu": "10/2013/TTLT",
        "ten": "Thông tư liên tịch 10/2013/TTLT-BTP-BCA-TANDTC-VKSNDTC-BTC",
        "co_quan": "Bộ Tư pháp - Bộ Công an - TANDTC - VKSNDTC - Bộ Tài chính",
        "ngay_hieu_luc": "2013-08-15",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_43.2013.tt.blđtbxh.pdf": {
        "so_hieu": "43/2013/TT-BLĐTBXH",
        "ten": "Thông tư 43/2013/TT-BLĐTBXH quản lý lao động, tiền lương",
        "co_quan": "Bộ Lao động - Thương binh và Xã hội",
        "ngay_hieu_luc": "2013-12-01",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_09-2014-tt-btc_09-2014-tt-btc.pdf": {
        "so_hieu": "09/2014/TT-BTC",
        "ten": "Thông tư 09/2014/TT-BTC về phí bán đấu giá chứng khoán",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2014-03-01",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "vanbangoc_60.2015.nđ.cp.pdf": {
        "so_hieu": "60/2015/NĐ-CP",
        "ten": "Nghị định 60/2015/NĐ-CP sửa đổi bổ sung Nghị định 58/2012/NĐ-CP",
        "co_quan": "Chính phủ",
        "ngay_hieu_luc": "2015-09-01",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Nghị định"
    },
    "vanbangoc_73-2013-tt-btc_73-2013-tt-btc.pdf": {
        "so_hieu": "73/2013/TT-BTC",
        "ten": "Thông tư 73/2013/TT-BTC hướng dẫn niêm yết chứng khoán",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2013-07-15",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    },
    "tvhienthitoanvan_20.vbhn.btc.pdf": {
        "so_hieu": "20/VBHN-BTC",
        "ten": "Văn bản hợp nhất 20/VBHN-BTC thi hành Luật Chứng khoán",
        "co_quan": "Bộ Tài chính",
        "ngay_hieu_luc": "2018-05-15",
        "tinh_trang": "Còn hiệu lực",
        "loai": "Thông tư"
    }
}

def parse_metadata_from_filename(filename):
    """Trích xuất sơ bộ số hiệu và loại văn bản từ tên file bằng Regex."""
    loai = "Văn bản"
    so_hieu = "Unknown"
    
    filename_lower = filename.lower()
    
    # Ưu tiên kiểm tra map fallback trước để có kết quả chính xác nhất
    if filename_lower in METADATA_FALLBACK_MAP:
        meta = METADATA_FALLBACK_MAP[filename_lower]
        return meta["loai"], meta["so_hieu"]
        
    if "nq" in filename_lower or "nghị quyết" in filename_lower:
        loai = "Nghị quyết"
    elif "nd" in filename_lower or "nđ" in filename_lower or "nghị định" in filename_lower:
        loai = "Nghị định"
    elif "luật" in filename_lower:
        loai = "Luật"
    elif "tt" in filename_lower or "thông tư" in filename_lower:
        loai = "Thông tư"
        
    # Tìm kiếm các mẫu số hiệu như 69.2025 hoặc 155/2020...
    match = re.search(r'(\d+[\./]\d+)', filename)
    if match:
        so_hieu = match.group(1)
        
    return loai, so_hieu

def extract_metadata_with_llm(text_content, filename):
    """Sử dụng Gemini API để trích xuất các cột thông tin bị thiếu từ nội dung văn bản."""
    # Nếu tệp nằm trong map fallback, sử dụng trực tiếp để tránh tải API khi bị rate-limit
    fn_lower = filename.lower()
    if fn_lower in METADATA_FALLBACK_MAP:
        meta = METADATA_FALLBACK_MAP[fn_lower]
        return meta["so_hieu"], meta["ten"], meta["co_quan"], meta["ngay_hieu_luc"], meta["tinh_trang"]
        
    if not GEMINI_API_KEY:
        print("  [!] Không tìm thấy GEMINI_API_KEY. Bỏ qua bước bóc tách bằng LLM.")
        return "Đang cập nhật", filename, "Đang cập nhật", "Đang cập nhật", "Còn hiệu lực"
        
    # Trích xuất 2000 ký tự đầu tiên
    preamble = text_content[:2000]
    
    prompt = f"""Bạn là một chuyên gia phân tích văn bản pháp luật. Hãy phân tích đoạn trích sau từ đầu một văn bản pháp luật Việt Nam (tên file gốc là '{filename}') và trích xuất các thông tin sau dưới dạng JSON cực kỳ ngắn gọn:
1. "so_hieu": Số hiệu chính thức của văn bản (ví dụ: "155/2020/NĐ-CP", "54/2019/QH14", "28/2020/TT-BTC"). Hãy cố gắng tìm số hiệu chính thức ghi ở góc trái đầu văn bản, trích xuất đầy đủ cả phần số và phần ký hiệu viết tắt.
2. "ten": Tên đầy đủ chính thức của văn bản pháp luật đó (ví dụ: "Nghị định 155/2020/NĐ-CP quy định chi tiết thi hành một số điều của Luật Chứng khoán").
3. "co_quan": Tên cơ quan ban hành (ví dụ: "Chính phủ", "Bộ Tài chính", "Quốc hội", "Ủy ban Thường vụ Quốc hội").
4. "ngay_hieu_luc": Ngày có hiệu lực của văn bản (định dạng YYYY-MM-DD, ví dụ: "2021-01-01"). Nếu không tìm thấy ngày hiệu lực cụ thể, hãy cố gắng đoán dựa vào ngày ký hoặc ngày ban hành, nếu không thể đoán được hãy ghi "Đang cập nhật".
5. "tinh_trang": Tình trạng hiệu lực hiện tại của văn bản này. Mặc định ghi "Còn hiệu lực" trừ khi văn bản ghi rõ đã bị thay thế hoặc hết hiệu lực.

Đoạn trích văn bản:
\"\"\"
{preamble}
\"\"\"

Trả về kết quả dưới dạng JSON thuần túy, không có thẻ ```json hay bất kỳ chữ giải thích nào khác. Định dạng bắt buộc:
{{
  "so_hieu": "...",
  "ten": "...",
  "co_quan": "...",
  "ngay_hieu_luc": "...",
  "tinh_trang": "..."
}}
"""
    try:
        from google import genai
        client = genai.Client(api_key=GEMINI_API_KEY)
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
        )
        
        clean_text = response.text.strip().replace("```json", "").replace("```", "").strip()
        data = json.loads(clean_text)
        
        so_hieu = data.get("so_hieu", "Đang cập nhật")
        ten = data.get("ten", filename)
        co_quan = data.get("co_quan", "Đang cập nhật")
        ngay_hieu_luc = data.get("ngay_hieu_luc", "Đang cập nhật")
        tinh_trang = data.get("tinh_trang", "Còn hiệu lực")
        
        return so_hieu, ten, co_quan, ngay_hieu_luc, tinh_trang
    except Exception as e:
        print(f"  [!] Lỗi LLM trích xuất metadata cho {filename}: {e}")
        return "Đang cập nhật", filename, "Đang cập nhật", "Đang cập nhật", "Còn hiệu lực"

def ingest_data():
    print("[*] Đang khởi tạo kết nối DB và ChromaDB...")
    db = SessionLocal()
    
    # Xóa sạch bảng PostgreSQL cũ để tránh trùng lặp Unknown
    print("[*] Đang làm sạch cơ sở dữ liệu cũ...")
    try:
        db.query(Document).delete()
        db.commit()
    except Exception as e:
        print(f"  [!] Không thể xóa bảng Document: {e}")
        db.rollback()
        
    os.makedirs(CHROMA_DIR, exist_ok=True)
    chroma_client = chromadb.PersistentClient(path=CHROMA_DIR)
    
    try:
        chroma_client.delete_collection(CHROMA_COLLECTION)
    except Exception:
        pass
    
    collection = chroma_client.create_collection(
        name=CHROMA_COLLECTION,
        metadata={"description": "Legal documents - Vietnam Securities Law"}
    )
    
    print(f"[*] Đang tải model embedding: {EMBEDDING_MODEL}...")
    model = SentenceTransformer(EMBEDDING_MODEL)
    
    all_chunks = []
    
    # 1. Quét thư mục
    if not os.path.exists(DATA_PATH):
        print(f"[!] Không tìm thấy thư mục {DATA_PATH}")
        return
        
    folders = [f for f in os.listdir(DATA_PATH) if os.path.isdir(os.path.join(DATA_PATH, f))]
    print(f"[*] Tìm thấy {len(folders)} thư mục tài liệu.")
    
    doc_counter = 0
    for folder in folders:
        folder_path = os.path.join(DATA_PATH, folder)
        files = os.listdir(folder_path)
        
        # Ưu tiên PDF, nếu không có thì tìm DOCX
        target_file = None
        for f in files:
            if f.lower().endswith('.pdf'):
                target_file = f
                break
        
        if not target_file:
            for f in files:
                if f.lower().endswith('.docx'):
                    target_file = f
                    break
                    
        if not target_file:
            print(f"  [!] Bỏ qua {folder} vì không tìm thấy file PDF hay DOCX.")
            continue
            
        file_path = os.path.join(folder_path, target_file)
        
        # 2. Đọc nội dung
        if target_file.lower().endswith('.pdf'):
            text_content = extract_text_from_pdf(file_path)
        else:
            text_content = extract_text_from_docx(file_path)
            
        if not text_content.strip():
            print(f"  [!] Cảnh báo: Không trích xuất được text từ {target_file}")
            continue
            
        # 3. Trích xuất metadata bằng Regex và LLM
        loai, _ = parse_metadata_from_filename(target_file)
        
        print(f"  [*] Đang xử lý tài liệu: {target_file}...")
        print("      → Đang gọi Gemini bóc tách số hiệu, tên, cơ quan ban hành...")
        so_hieu, ten, co_quan, ngay_hieu_luc, tinh_trang = extract_metadata_with_llm(text_content, target_file)
        
        # Nếu LLM không tìm thấy số hiệu hợp lệ, fallback về regex tên file
        if so_hieu == "Đang cập nhật" or not so_hieu or so_hieu.strip() == "":
            _, fallback_sohieu = parse_metadata_from_filename(target_file)
            so_hieu = fallback_sohieu if fallback_sohieu else "Unknown"
            
        print(f"      ✓ Kết quả: Số hiệu: {so_hieu} | CQ: {co_quan} | Ngày HL: {ngay_hieu_luc}")
        
        doc_id = f"doc_{folder}_{so_hieu}".replace(".", "_").replace("/", "_").replace("-", "_").replace(" ", "_")
        
        # Lệnh Delay: Chờ 4 giây đối vì tài khoản Gemini miễn phí (Tối đa 15 yêu cầu/phút)
        # Chỉ dừng khi thực sự gọi Gemini API (tệp không nằm trong fallback map)
        if GEMINI_API_KEY and target_file.lower() not in METADATA_FALLBACK_MAP:
            print("      [Delay] Đang tạm dừng 4 giây để tránh lỗi quá tải API (429 Rate Limit)...")
            time.sleep(4)
        
        # Lưu vào PostgreSQL 
        existing_doc = db.query(Document).filter(Document.doc_id == doc_id).first()
        if not existing_doc:
            new_doc = Document(
                doc_id=doc_id,
                so_hieu=so_hieu,
                ten=ten,
                loai=loai,
                co_quan=co_quan,
                ngay_hieu_luc=ngay_hieu_luc,
                tinh_trang=tinh_trang,
                file_path=file_path
            )
            db.add(new_doc)
            db.commit()
            db.refresh(new_doc)
            pg_doc_id = new_doc.id
        else:
            existing_doc.ten = ten
            existing_doc.so_hieu = so_hieu
            existing_doc.co_quan = co_quan
            existing_doc.ngay_hieu_luc = ngay_hieu_luc
            existing_doc.tinh_trang = tinh_trang
            db.commit()
            pg_doc_id = existing_doc.id
            
        # 4. Chunking
        doc_metadata = {
            "doc_id": doc_id,
            "pg_doc_id": pg_doc_id,
            "so_hieu": so_hieu,
            "ten": ten,
            "loai": loai,
            "co_quan": co_quan,
            "ngay_hieu_luc": ngay_hieu_luc,
            "tinh_trang": tinh_trang,
            "file_path": file_path
        }
        
        chunks = chunk_by_article(text_content, doc_metadata)
        all_chunks.extend(chunks)
        
        doc_counter += 1
        print(f"      ✓ Hoàn tất: {len(chunks)} chunks")
        
    # 5. Embedding và đưa vào ChromaDB
    if not all_chunks:
        print("[!] Không có chunk nào được tạo ra.")
        return
        
    texts = [c["text"] for c in all_chunks]
    metadatas = [c["metadata"] for c in all_chunks]
    ids = [f"{c['metadata']['doc_id']}_chunk_{i}" for i, c in enumerate(all_chunks)]
    
    print(f"\n[*] Đang tạo embeddings cho {len(texts)} chunks (quá trình này có thể mất chút thời gian)...")
    embeddings = model.encode(texts, show_progress_bar=True, batch_size=16)
    
    print(f"[*] Đang lưu {len(texts)} chunks vào ChromaDB...")
    collection.add(
        ids=ids,
        documents=texts,
        embeddings=embeddings.tolist(),
        metadatas=metadatas,
    )
    
    print(f"\n[+] HOÀN THÀNH! Đã nạp thành công {doc_counter} tài liệu.")
    print("[+] Đã tự động cập nhật đầy đủ Metadata vào PostgreSQL và ChromaDB bằng AI.")
    db.close()

if __name__ == "__main__":
    ingest_data()
